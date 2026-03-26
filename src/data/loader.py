from typing import List, Optional, Dict, Any, Union
from pathlib import Path
from langchain_community.document_loaders import PyPDFLoader, UnstructuredPDFLoader, UnstructuredWordDocumentLoader
from langchain_core.documents import Document as LangChainDocument
from src.utils.logger import logger
from src.config.settings import settings
import docx


class DocumentLoader:
    def __init__(self, loader_type: str = "unstructured"):
        self.loader_type = loader_type
        logger.info(f"Initialized DocumentLoader with {loader_type} backend")

    def load_document(self, file_path: Path) -> List[LangChainDocument]:
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

        extension = file_path.suffix.lower()

        if extension == ".pdf":
            return self._load_pdf(file_path)
        elif extension == ".docx":
            return self._load_docx(file_path)
        elif extension == ".txt":
            return self._load_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")

    def _load_pdf(self, file_path: Path) -> List[LangChainDocument]:
        logger.info(f"Loading PDF from {file_path}")

        try:
            if self.loader_type == "pypdf":
                loader = PyPDFLoader(str(file_path))
            elif self.loader_type == "unstructured":
                loader = UnstructuredPDFLoader(
                    str(file_path),
                    mode="elements",
                    strategy="fast"
                )
            else:
                loader = PyPDFLoader(str(file_path))

            documents = loader.load()
            logger.info(f"Successfully loaded {len(documents)} pages/elements from PDF")

            for doc in documents:
                doc.metadata["source"] = str(file_path.name)
                doc.metadata["format"] = "pdf"

            return documents

        except Exception as e:
            logger.error(f"Failed to load PDF: {e}")
            raise

    def _load_docx(self, file_path: Path) -> List[LangChainDocument]:
        logger.info(f"Loading DOCX from {file_path}")

        try:
            loader = UnstructuredWordDocumentLoader(
                str(file_path),
                mode="elements"
            )

            documents = loader.load()
            if not documents:
                logger.warning("Unstructured loader returned empty, falling back to python-docx")
                documents = self._load_docx_fallback(file_path)

            logger.info(f"Successfully loaded {len(documents)} elements from DOCX")

            for doc in documents:
                doc.metadata["source"] = str(file_path.name)
                doc.metadata["format"] = "docx"

            return documents

        except Exception as e:
            logger.error(f"Failed to load DOCX with Unstructured: {e}")
            return self._load_docx_fallback(file_path)

    def _load_docx_fallback(self, file_path: Path) -> List[LangChainDocument]:
        try:
            doc = docx.Document(file_path)
            documents = []
            current_section = None

            for para in doc.paragraphs:
                if not para.text.strip():
                    continue

                is_heading = para.style.name.startswith('Heading') or para.style.name.startswith('Заголовок')
                doc_obj = LangChainDocument(
                    page_content=para.text,
                    metadata={
                        "source": str(file_path.name),
                        "format": "docx",
                        "paragraph_index": len(documents),
                        "is_heading": is_heading,
                        "style": para.style.name,
                        "section": current_section
                    }
                )
                documents.append(doc_obj)

                if is_heading:
                    current_section = para.text

            logger.info(f"Loaded {len(documents)} paragraphs from DOCX (fallback)")
            return documents

        except Exception as e:
            logger.error(f"Failed to load DOCX with fallback: {e}")
            raise

    def _load_txt(self, file_path: Path) -> List[LangChainDocument]:
        logger.info(f"Loading TXT from {file_path}")

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            document = LangChainDocument(
                page_content=content,
                metadata={
                    "source": str(file_path.name),
                    "format": "txt"
                }
            )

            logger.info("Successfully loaded TXT file")
            return [document]

        except Exception as e:
            logger.error(f"Failed to load TXT: {e}")
            raise

    def load_multiple_documents(self, directory: Path, pattern: str = "*.*") -> List[LangChainDocument]:
        all_documents = []
        supported_extensions = [".pdf", ".docx", ".txt"]

        for ext in supported_extensions:
            for doc_path in directory.glob(f"*{ext}"):
                try:
                    documents = self.load_document(doc_path)
                    all_documents.extend(documents)
                    logger.info(f"Loaded {len(documents)} from {doc_path.name}")
                except Exception as e:
                    logger.error(f"Failed to load {doc_path.name}: {e}")
                    continue

        return all_documents


_loader_instance: Optional[DocumentLoader] = None


def get_document_loader(loader_type: str = "unstructured") -> DocumentLoader:
    global _loader_instance
    if _loader_instance is None:
        _loader_instance = DocumentLoader(loader_type)
    return _loader_instance
